import json

import ollama
from openai import OpenAI
import streamlit as st
from database.db import (
    create_user,
    login_user
)

from database.db import (
    save_chat,
    get_all_chats,
    get_chat_by_id,
    update_chat,
    delete_chat,
    get_chat_titles,
    get_latest_chat,
    get_total_chats,
    get_max_chat_id,
    get_rag_chats,
    get_all_rag_chats,
    search_chats,


)

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from database.vector_db import (
    save_chunk,
    get_documents,
    get_document_count,
    delete_document

)

from database.vector_search import (
    search_chunks
)

from utils.chunker import (
    chunk_text
)

from utils.embeddings import (
    create_embedding
)

from utils.file_reader import (
    read_pdf,
    read_docx,
    read_txt
)

from datetime import datetime

# ------------------------------------------------
# PAGE CONFIG
# ------------------------------------------------

st.set_page_config(
    page_title="OllamaCopilot",
    page_icon="🤖",
    layout="wide"
)

# ------------------------------------------------
# SESSION STATE
# ------------------------------------------------
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if "username" not in st.session_state:
    st.session_state.username = None

if "userid" not in st.session_state:
    st.session_state.userid = None
if "role" not in st.session_state:

    st.session_state.role = "user"

DEFAULTS = {

    # CHAT

    "messages": [],
    "chat_title": "New Conversation",
    "current_chat_id": None,
    "logged_in": "",
    "username": "",
    "userid":"",
    "role":"",

    # RAG

    "document_text": "",
    "selected_document": None,
"rag_messages": [],
"current_rag_chat_id": None,

    # APP

    "app_mode": "💬 Chat"
}

for key, value in DEFAULTS.items():

    if key not in st.session_state:

        st.session_state[key] = value

# ------------------------------------------------
# MODELS
# ------------------------------------------------
if not st.session_state.logged_in :
 st.sidebar.subheader(
    "🔐 Authentication"
)
 auth_mode = st.sidebar.radio(
    "Select",
    [
        "Login",
        "Signup"
    ]
)
 if auth_mode == "Signup":

    username = st.sidebar.text_input(
        "Username"
    )

    password = st.sidebar.text_input(
        "Password",
        type="password"
    )

    if st.sidebar.button(
        "Create Account"
    ):

        try:

            create_user(
                username,
                password
            )

            st.sidebar.success(
                "Account Created"
            )

        except Exception as e:

            st.sidebar.error(
                str(e)
            )
 if auth_mode == "Login":

    username = st.sidebar.text_input(
        "Username"
    )

    password = st.sidebar.text_input(
        "Password",
        type="password"
    )

    if st.sidebar.button(
        "Login"
    ):

        user = login_user(
            username,
            password
        )

        if user:

            st.session_state.logged_in = True

            st.session_state.userid = (
                user[0]
            )

            st.session_state.username = (
                user[1]
            )

            st.session_state.role = (
                user[4]
            )

            if "role" not in st.session_state:
                st.session_state.role = "user"

            st.rerun()

        else:

            st.sidebar.error(
                "Invalid Login"
            )

else:

 if st.session_state.logged_in:

     st.sidebar.success(
         "✅ Hybrid Search Enabled"
     )

     st.sidebar.success(
         "✅ Re-ranking Enabled"
     )

     st.sidebar.success(
         f"👤 {st.session_state.username}"
     )

     if st.sidebar.button(
             "🚪 Logout"
     ):
         st.session_state.logged_in = False

         st.session_state.username = None

         st.session_state.role = None

         st.session_state.messages = []

         st.session_state.rag_messages = []

         st.session_state.current_chat_id = None

         st.session_state.current_rag_chat_id = None

         st.session_state.chat_title = (
             "New Conversation"
         )
         if st.session_state.logged_in:

             st.sidebar.success(
                 f"👤 {st.session_state.username}"
             )

             if st.sidebar.button(
                     "🚪 Logout"
             ):
                 st.session_state.logged_in = False

                 st.session_state.username = None

                 st.session_state.role = None

                 st.session_state.messages = []

                 st.session_state.rag_messages = []

                 st.session_state.current_chat_id = None

                 st.session_state.current_rag_chat_id = None

                 st.session_state.chat_title = (
                     "New Conversation"
                 )
                 st.rerun()

 provider = st.sidebar.selectbox(
    "Provider",
    [
        "Ollama",
        "OpenAI"
    ]
)

 openai_api_key = ""

 if provider == "OpenAI":

    openai_api_key = st.sidebar.text_input(
        "OpenAI API Key",
        type="password"
    )

 if provider == "Ollama":

    MODELS = [

        "phi3:latest",

        "llama3:8b",

        "llama3.1:8b",

        "qwen3:8b",

        "gemma3:4b",

        "mistral:latest",

        "deepseek-coder:latest"
    ]

 else:

    MODELS = [

        "gpt-4o-mini",

        "gpt-4.1-mini",

        "gpt-4.1"
    ]
 def get_stream_response(
        provider,
        model,
        messages,
        api_key=""
):

    response_placeholder = st.empty()

    full_response = ""

    if provider == "Ollama":

        stream = ollama.chat(
            model=model,
            messages=messages,
            stream=True
        )

        for chunk in stream:

            token = chunk[
                "message"
            ]["content"]

            full_response += token

            response_placeholder.markdown(
                full_response + "▌"
            )

    else:

        client = OpenAI(
            api_key=api_key
        )

        stream = (
            client.chat.completions.create(
                model=model,
                messages=messages,
                stream=True
            )
        )

        for chunk in stream:

            if (
                chunk.choices[0]
                .delta.content
            ):

                token = (
                    chunk.choices[0]
                    .delta.content
                )

                full_response += token

                response_placeholder.markdown(
                    full_response + "▌"
                )

    response_placeholder.markdown(
        full_response
    )

    return full_response

 def get_ai_response(
        provider,
        model,
        messages,
        api_key=""
):

    if provider == "Ollama":

        response = ollama.chat(
            model=model,
            messages=messages
        )

        return response[
            "message"
        ]["content"]

    else:

        client = OpenAI(
            api_key=api_key
        )

        response = (
            client.chat.completions.create(
                model=model,
                messages=messages
            )
        )

        return (
            response
            .choices[0]
            .message.content
        )

# ------------------------------------------------
# SIDEBAR
# ------------------------------------------------
 def export_pdf():

    pdf_file = "chat_export.pdf"

    doc = SimpleDocTemplate(
        pdf_file
    )

    styles = getSampleStyleSheet()

    content = [Paragraph(
        "OllamaCopilot Chat Export",
        styles["Title"]
    ), Spacer(1, 12)]

    for msg in (
        st.session_state.messages
    ):

        content.append(
            Paragraph(
                f"<b>{msg['role'].upper()}</b>",
                styles["Heading2"]
            )
        )

        content.append(
            Paragraph(
                msg["content"],
                styles["BodyText"]
            )
        )

        content.append(
            Spacer(1,10)
        )

    doc.build(content)

    return pdf_file
 def export_docx():

    doc = Document()

    doc.add_heading(
        "OllamaCopilot Chat Export",
        0
    )

    for msg in (
        st.session_state.messages
    ):

        doc.add_heading(
            msg["role"].upper(),
            level=1
        )

        doc.add_paragraph(
            msg["content"]
        )

    file_name = (
        "chat_export.docx"
    )

    doc.save(
        file_name
    )

    return file_name

 st.sidebar.title(
    "🤖 OllamaCopilot"
)

# ------------------------------------------------
# MODE
# ------------------------------------------------

 st.session_state.app_mode = (
    st.sidebar.radio(
        "Select Mode",
        [
            "💬 Chat",
            "📄 RAG"
        ]
    )
)

# ------------------------------------------------
# MODEL
# ------------------------------------------------

 model = st.sidebar.selectbox(
    "Select Model",
    MODELS,
    index=0
)

# ------------------------------------------------
# STATISTICS
# ------------------------------------------------
 st.sidebar.subheader(
    "📊 Statistics"
)

 try:

    total_chats = (
        get_total_chats(st.session_state.userid)
    )

    total_docs = (
        get_document_count(st.session_state.userid)
    )

    latest_chat_id = (
        get_max_chat_id()
    )
    col3,col4 = (
        st.sidebar.columns(2)
    )

    with col3:
      st.metric(
        "Latest Chat ID",
        latest_chat_id
       )

    col1, col2 = (
        st.sidebar.columns(2)
    )

    with col1:

        st.metric(
            "Chats",
            total_chats
        )

    with col2:

        st.metric(
            "Docs",
            total_docs
        )


 except Exception as e:

    st.sidebar.error(
        f"Statistics Error: {e}"
    )
#------------------------------
#----------Search--------------
#------------------------------
 st.sidebar.subheader(
    "🔍 Search Chats"
)

 search_text = st.sidebar.text_input(
    "Search Title"
)

 if search_text:

    results = search_chats(
        search_text,
        st.session_state.userid
    )

    st.sidebar.write(
        f"Found: {len(results)}"
    )

    for chat in results:

        if st.sidebar.button(
            f"{chat[0]} - {chat[1]}",
            key=f"search_{chat[0]}"
        ):

            record = get_chat_by_id(
                chat[0]
            )

            if record:

                st.session_state.current_chat_id = (
                    record[0]
                )

                st.session_state.chat_title = (
                    record[1]
                )

                try:

                    st.session_state.messages = (
                        json.loads(
                            record[4]
                        )
                    )

                except:

                    pass

                st.rerun()


# ------------------------------------------------
# MODE SPECIFIC SIDEBAR
# ------------------------------------------------

 if (
    st.session_state.app_mode
    ==
    "💬 Chat"
):

    st.sidebar.success(
        "Chat Mode Active"
    )


 else:

    st.sidebar.info(
        "RAG Mode Active"
    )

 chat_export = f"""
====================================
OllamaCopilot Chat Export
====================================

Chat Title:
{st.session_state.chat_title}

Export Date:
{datetime.now()}

Total Messages:
{len(st.session_state.messages)}

====================================
Conversation
====================================

"""

 for msg in st.session_state.messages:

    role = msg["role"].upper()

    chat_export += f"""

------------------------------------
{role}
------------------------------------

{msg['content']}

"""

 st.sidebar.download_button(
    "⬇ Export Chat",
    chat_export,
    file_name=f"{st.session_state.chat_title}.txt",
    mime="text/plain"
)

 if st.sidebar.button(
    "📄 Export PDF"
):

    pdf_file = export_pdf()

    with open(
        pdf_file,
        "rb"
    ) as f:

        st.sidebar.download_button(
            "Download PDF",
            f,
            file_name=
            "chat_export.pdf"
        )


 if st.sidebar.button(
    "📝 Export DOCX"
):

    docx_file = (
        export_docx()
    )

    with open(
        docx_file,
        "rb"
    ) as f:

        st.sidebar.download_button(
            "Download DOCX",
            f,
            file_name=
            "chat_export.docx"
        )

 chat_json = json.dumps(
    st.session_state.messages,
    indent=4,
    ensure_ascii=False
)

 st.sidebar.download_button(
    "📦 Export JSON",
    chat_json,
    file_name=
    "chat_export.json",
    mime=
    "application/json"
)




# ------------------------------------------------
# CHAT SIDEBAR
# ------------------------------------------------
 if (
    st.session_state.app_mode
    ==
    "💬 Chat"
):

    # --------------------------------------------
    # LATEST CHAT
    # --------------------------------------------

    st.sidebar.subheader(
        "💬 Latest Chat"
    )

    try:

        latest = get_latest_chat()

        if latest:

            st.sidebar.success(
                latest[1]
            )

        else:

            st.sidebar.info(
                "No Chats Yet"
            )

    except Exception as e:

        st.sidebar.error(
            str(e)
        )

    # --------------------------------------------
    # CURRENT CHAT
    # --------------------------------------------

    st.sidebar.subheader(
        "📝 Current Chat"
    )

    st.sidebar.write(
        st.session_state.chat_title
    )

    # --------------------------------------------
    # DOWNLOAD CHAT
    # --------------------------------------------

    chat_text = ""

    for msg in (
        st.session_state.messages
    ):

        chat_text += (
            f"{msg['role']} : "
            f"{msg['content']}\n\n"
        )

    st.sidebar.download_button(
        label="⬇ Download Chat",
        data=chat_text,
        file_name="conversation.txt",
        mime="text/plain"
    )




    # --------------------------------------------
    # SAVED CHATS
    # --------------------------------------------

    st.sidebar.subheader(
        "📚 Saved Chats"
    )

    try:

        chats = get_all_chats(st.session_state.userid)

        if chats:

            for chat in chats:

                c1, c2, c3 = (
                    st.sidebar.columns(
                        [4, 1, 1]
                    )
                )

                # ----------------------
                # LOAD CHAT
                # ----------------------

                with c1:

                    if st.button(
                        f"{chat[0]} - {chat[1]}",
                        key=f"load_{chat[0]}"
                    ):

                        record = (
                            get_chat_by_id(
                                chat[0]
                            )
                        )

                        if record:

                            st.session_state.current_chat_id = (
                                record[0]
                            )

                            st.session_state.chat_title = (
                                record[1]
                            )

                            try:

                                st.session_state.messages = (
                                    json.loads(
                                        record[4]
                                    )
                                )

                            except:

                                st.session_state.messages = []

                            st.rerun()

                # ----------------------
                # UPDATE CHAT
                # ----------------------

                with c2:

                    if st.button(
                        "💾",
                        key=f"update_{chat[0]}"
                    ):

                        try:

                            chat_text = ""

                            for msg in (
                                st.session_state.messages
                            ):

                                chat_text += (
                                    f"{msg['role']} : "
                                    f"{msg['content']}\n\n"
                                )

                            messages_json = (
                                json.dumps(
                                    st.session_state.messages
                                )
                            )

                            update_chat(
                                chat[0],
                                st.session_state.chat_title,
                                chat_text,
                                messages_json,
                                st.session_state.userid
                            )

                            st.sidebar.success(
                                "Updated"
                            )

                        except Exception as e:

                            st.sidebar.error(
                                str(e)
                            )

                # ----------------------
                # DELETE CHAT
                # ----------------------

                with c3:

                    if st.button(
                        "❌",
                        key=f"delete_{chat[0]}"
                    ):

                        try:

                            delete_chat(
                                chat[0],
                                st.session_state.userid
                            )

                            st.rerun()

                        except Exception as e:

                            st.sidebar.error(
                                str(e)
                            )

        else:

            st.sidebar.info(
                "No Saved Chats"
            )

    except Exception as e:

        st.sidebar.error(
            f"Chat Error: {e}"
        )

    # --------------------------------------------
    # HISTORY
    # --------------------------------------------

    st.sidebar.subheader(
        "📝 History"
    )

    try:

        history = get_chat_titles(st.session_state.userid)

        if history:

            for chat in history:

                if st.sidebar.button(
                        f"{chat[0]} - {chat[1]}",
                        key=f"history_{chat[0]}"
                ):

                    record = get_chat_by_id(
                        chat[0]
                    )

                    if record:

                        st.session_state.current_chat_id = (
                            record[0]
                        )

                        st.session_state.chat_title = (
                            record[1]
                        )

                        try:

                            st.session_state.messages = (
                                json.loads(
                                    record[4]
                                )
                            )

                        except Exception:

                            st.session_state.messages = []

                        st.rerun()

        else:

            st.sidebar.info(
                "No History Found"
            )

    except Exception as e:

        st.sidebar.error(
            str(e)
        )

    # --------------------------------------------
    # NEW CHAT
    # --------------------------------------------

    if st.sidebar.button(
        "🆕 New Conversation",
        use_container_width=True
    ):

        st.session_state.messages = []

        st.session_state.chat_title = (
            "New Conversation"
        )

        st.session_state.current_chat_id = (
            None
        )

        st.rerun()


# ------------------------------------------------
# RAG SIDEBAR
# ------------------------------------------------

 if (
    st.session_state.app_mode
    ==
    "📄 RAG"
):

    st.sidebar.subheader(
        "📄 RAG Documents"
    )

    uploaded_file = (
        st.sidebar.file_uploader(
            "Upload File",
            type=[
                "pdf",
                "docx",
                "txt"
            ]
        )
    )
    # --------------------------------------------
    # SAVED RAG CHATS
    # --------------------------------------------

    st.sidebar.subheader(
        "📄 Saved RAG Chats"
    )

    try:

        rag_chats = get_all_rag_chats(st.session_state.userid)

        if rag_chats:

            for chat in rag_chats:

                c1, c2, c3 = st.sidebar.columns(
                    [4, 1, 1]
                )

                # -------------------------
                # LOAD
                # -------------------------

                with c1:

                    if st.button(
                            chat[1],
                            key=f"rag_load_{chat[0]}"
                    ):

                        record = get_chat_by_id(
                            chat[0]
                        )

                        if record:

                            st.session_state.current_rag_chat_id = (
                                record[0]
                            )

                            try:

                                st.session_state.rag_messages = (
                                    json.loads(
                                        record[4]
                                    )
                                )

                            except:

                                st.session_state.rag_messages = []

                            st.rerun()

                # -------------------------
                # UPDATE
                # -------------------------

                with c2:

                    if st.button(
                            "💾",
                            key=f"rag_update_{chat[0]}"
                    ):

                        try:

                            rag_chat_text = ""

                            for msg in (
                                    st.session_state.rag_messages
                            ):
                                rag_chat_text += (
                                    f"{msg['role']} : "
                                    f"{msg['content']}\n\n"
                                )

                            messages_json = (
                                json.dumps(
                                    st.session_state.rag_messages,
                                    ensure_ascii=False
                                )
                            )

                            update_chat(
                                chat[0],
                                chat[1],
                                rag_chat_text,
                                messages_json,
                                st.session_state.userid
                            )

                            st.sidebar.success(
                                "RAG Updated"
                            )

                        except Exception as e:

                            st.sidebar.error(
                                str(e)
                            )

                # -------------------------
                # DELETE
                # -------------------------

                with c3:

                    if st.button(
                            "❌",
                            key=f"rag_delete_{chat[0]}"
                    ):
                        delete_chat(
                            chat[0],
                            st.session_state.userid
                        )

                        st.rerun()

        else:

            st.sidebar.info(
                "No RAG Chats"
            )

    except Exception as e:

        st.sidebar.error(
            str(e)
        )

    # --------------------------------------------
    # DOCUMENT PROCESSING
    # --------------------------------------------

    if uploaded_file:

        text = ""

        try:

            file_name = (
                uploaded_file.name.lower()
            )

            if file_name.endswith(
                ".pdf"
            ):

                text = read_pdf(
                    uploaded_file
                )

            elif file_name.endswith(
                ".docx"
            ):

                text = read_docx(
                    uploaded_file
                )

            elif file_name.endswith(
                ".txt"
            ):

                text = read_txt(
                    uploaded_file
                )

            st.session_state.document_text = (
                text
            )

            st.success(
                f"✅ Loaded: {uploaded_file.name}"
            )

        except Exception as e:

            st.error(
                f"Document Error: {e}"
            )

    # --------------------------------------------
    # DOCUMENT PREVIEW
    # --------------------------------------------

    if (
        st.session_state.document_text
        and
        st.session_state.document_text.strip()
    ):

        st.subheader(
            "📄 Document Preview"
        )

        preview_length = min(
            5000,
            len(
                st.session_state.document_text
            )
        )

        st.text_area(
            "Content",
            st.session_state.document_text[
                :preview_length
            ],
            height=250
        )

        # ----------------------------------------
        # CREATE CHUNKS
        # ----------------------------------------

        try:

            chunks = chunk_text(
                st.session_state.document_text
            )

            st.info(
                f"📦 Chunks Created: {len(chunks)}"
            )

        except Exception as e:

            chunks = []

            st.error(
                f"Chunk Error: {e}"
            )

        # ----------------------------------------
        # INDEX DOCUMENT
        # ----------------------------------------

        if st.button(
            "🚀 Index Document",
            use_container_width=True
        ):

            if not chunks:

                st.warning(
                    "No chunks generated."
                )

            else:

                total_chunks = len(
                    chunks
                )

                progress_bar = (
                    st.progress(0)
                )

                status = st.empty()

                success_count = 0

                try:

                    for index, chunk in enumerate(
                        chunks
                    ):

                        status.write(
                            f"Processing "
                            f"{index + 1}"
                            f"/{total_chunks}"
                        )

                        embedding = (
                            create_embedding(
                                chunk
                            )
                        )

                        save_chunk(
                            uploaded_file.name,
                            chunk,
                            embedding,
                            st.session_state.userid
                        )

                        success_count += 1

                        progress_bar.progress(
                            (
                                index + 1
                            )
                            /
                            total_chunks
                        )

                    status.empty()

                    st.success(
                        f"✅ Indexed "
                        f"{success_count} chunks"
                    )
                    st.metric(
                        "Chunks Indexed",
                        success_count
                    )

                except Exception as e:

                    st.error(
                        f"Index Error: {e}"
                    )


    # --------------------------------------------
    # SAVED DOCUMENTS
    # --------------------------------------------

    st.sidebar.subheader(
        "📚 Saved Documents"
    )

    try:

        documents = get_documents(st.session_state.userid)

        for doc in documents:

            c1, c2 = st.sidebar.columns(
                [5, 1]
            )
            with c1:

             if st.button(
                    f"📄 {doc[0]}",
                    key=f"doc_{doc[0]}"
            ):
                st.session_state.selected_document = (
                    doc[0]
                )

                st.rerun()

            with c2:

                if st.button(
                        "❌",
                        key=f"doc_delete_{doc[0]}"
                ):
                    delete_document(
                        doc[0],
                        st.session_state.userid
                    )

                    st.success(
                        "Deleted"
                    )

                    st.rerun()

    except Exception as e:

        st.sidebar.error(
            f"Document Error: {e}"
        )

    st.sidebar.subheader(
        "📄 Current Document"
    )

    if st.session_state.selected_document:

        st.sidebar.success(
            st.session_state.selected_document
        )

    else:

        st.sidebar.warning(
            "No Document Selected"
        )


    # --------------------------------------------
    # RAG HISTORY
    # --------------------------------------------

    st.sidebar.subheader(
        "📄 RAG History"
    )

    try:

        rag_chats1 = get_rag_chats(st.session_state.userid)

        if rag_chats1:

            for chat in rag_chats1:

                c1, c2 = st.sidebar.columns(
                    [5, 1]
                )

                with c1:

                    if st.sidebar.button(
                            f"{chat[0]} - {chat[1]}",
                            key=f"history_load_{chat[0]}"
                    ):

                        record1 = get_chat_by_id(
                            chat[0]
                        )

                        if record1:

                            try:

                                st.session_state.rag_messages = (
                                    json.loads(
                                        record1[4]
                                    )
                                )

                                st.session_state.current_rag_chat_id = (
                                    record1[0]
                                )

                                st.rerun()

                            except Exception as e:

                                st.sidebar.error(
                                    str(e)
                                )


        else:

            st.sidebar.info(
                "No RAG History"
            )

    except Exception as e:

        st.sidebar.error(
            str(e)
        )



    # --------------------------------------------
    # DOCUMENT COUNT
    # --------------------------------------------

    st.sidebar.subheader(
        "📊 Documents"
    )

    try:

        st.sidebar.write(
            f"Total Documents: "
            f"{get_document_count(st.session_state.userid)}"
        )

    except Exception as e:

        st.sidebar.error(
            str(e)
        )

    if st.sidebar.button(
            "🆕 New RAG Chat"
    ):
        st.session_state.rag_messages = []

        st.session_state.current_rag_chat_id = None

        st.rerun()


# ------------------------------------------------
# MAIN LAYOUT
# ------------------------------------------------

if (
    st.session_state.app_mode
    ==
    "💬 Chat"
):

    st.title(
        "💬 Chat Assistant"
    )

    # Chat code here

else:

    st.title(
        "📄 RAG Assistant"
    )

    # RAG code here

# =================================================
# CHAT TAB
# =================================================

if st.session_state.app_mode == "💬 Chat":


    st.caption(
        "Normal Ollama Chat - No RAG"
    )

    # --------------------------------------------
    # DISPLAY CHAT HISTORY
    # --------------------------------------------

    for message in (
        st.session_state.messages
    ):

        with st.chat_message(
            message["role"]
        ):

            st.markdown(
                message["content"]
            )

    # --------------------------------------------
    # CHAT INPUT
    # --------------------------------------------

    prompt = st.chat_input(
        "Ask anything..."
    )

    if prompt:

        # ----------------------------------------
        # FIRST MESSAGE = TITLE
        # ----------------------------------------

        if len(
            st.session_state.messages
        ) == 0:

            st.session_state.chat_title = (
                prompt[:40]
            )

        # ----------------------------------------
        # USER MESSAGE
        # ----------------------------------------

        st.session_state.messages.append(
            {
                "role": "user",
                "content": prompt
            }
        )

        with st.chat_message(
            "user"
        ):

            st.markdown(
                prompt
            )

        # ----------------------------------------
        # BUILD MESSAGE LIST
        # ----------------------------------------

        final_messages = []

        final_messages.extend(
            st.session_state.messages[-10:]
        )

        # ----------------------------------------
        # OLLAMA RESPONSE
        # ----------------------------------------

        try:

            with st.chat_message(
                "assistant"
            ):

                response_placeholder = (
                    st.empty()
                )

                full_response = ""

                with st.spinner(
                    "Thinking..."
                ):
                    full_response = (
                        get_stream_response(
                            provider,
                            model,
                            final_messages,
                            openai_api_key
                        )
                    )


                response_placeholder.markdown(
                    full_response
                )

        except Exception as e:

            full_response = (
                f"❌ Ollama Error:\n\n{e}"
            )

            st.error(
                full_response
            )

        # ----------------------------------------
        # SAVE ASSISTANT MESSAGE
        # ----------------------------------------

        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": full_response
            }
        )

        # ----------------------------------------
        # BUILD CHAT TEXT
        # ----------------------------------------

        chat_text = ""

        for msg in (
            st.session_state.messages
        ):

            chat_text += (
                f"{msg['role']} : "
                f"{msg['content']}\n\n"
            )

        messages_json = (
            json.dumps(
                st.session_state.messages,
                ensure_ascii=False
            )
        )

        # ----------------------------------------
        # DATABASE SAVE
        # ----------------------------------------

        try:

            if st.session_state.current_chat_id is None:

                inserted_id = save_chat(
                    st.session_state.chat_title,
                    chat_text,
                    messages_json,
                    st.session_state.userid
                )

                st.session_state.current_chat_id = (
                    inserted_id
                )

            else:

                update_chat(
                    st.session_state.current_chat_id,
                    st.session_state.chat_title,
                    chat_text,
                    messages_json,
                    st.session_state.userid
                )

        except Exception as e:

            st.error(
                f"Database Error: {e}"
            )

        st.rerun()


# =================================================
# RAG TAB
# =================================================

else :

    st.caption(
        "Document Question Answering"
    )

    st.success(
        "RAG Mode Active"
    )

    if "rag_messages" in st.session_state:

        for message in st.session_state.rag_messages:
            with st.chat_message(
                    message["role"]
            ):
                st.markdown(
                    message["content"]
                )

    rag_prompt = st.chat_input(
        "Ask from documents..."
    )

    if rag_prompt:

        with st.chat_message(
            "user"
        ):

            st.markdown(
                rag_prompt
            )

        # ----------------------------------------
        # SEARCH DOCUMENTS
        # ----------------------------------------

        rag_context = ""

        try:

            results = search_chunks(
                question=rag_prompt,
                selected_document=st.session_state.selected_document,
                top_k=1,
                user_id=st.session_state.userid
            )

            for score, chunk in results:

                try:

                    chunk_text = (
                        chunk[2]
                    )


                    rag_context += (
                            chunk_text[:500]
                            + "\n\n"
                    )

                except Exception:
                    pass

        except Exception as e:

            st.error(
                f"Search Error: {e}"
            )

        # ----------------------------------------
        # NO DOCUMENT FOUND
        # ----------------------------------------

        if not rag_context.strip():

            st.warning(
                "No relevant documents found."
            )

        else:

            # ------------------------------------
            # BUILD PROMPT
            # ------------------------------------

            final_messages = [

                {
                    "role": "system",
                    "content": f"""
You are a RAG assistant.

Answer ONLY from the provided context.

If the answer is not present
say:

'I could not find that information
in the indexed documents.'

Context:

{rag_context}
"""
                },

                {
                    "role": "user",
                    "content": rag_prompt
                }
            ]

            # ------------------------------------
            # OLLAMA RESPONSE
            # ------------------------------------

            try:

                with st.chat_message(
                    "assistant"
                ):

                    response_placeholder = (
                        st.empty()
                    )

                    full_response = ""

                    with st.spinner(
                        "Searching Documents..."
                    ):

                        full_response = (
                            get_stream_response(
                                provider,
                                model,
                                final_messages,
                                openai_api_key
                            )
                        )

                    response_placeholder.markdown(
                        full_response
                    )
                    if "rag_messages" not in st.session_state:
                        st.session_state.rag_messages = []

                    st.session_state.rag_messages.append(
                        {
                            "role": "user",
                            "content": rag_prompt
                        }
                    )

                    st.session_state.rag_messages.append(
                        {
                            "role": "assistant",
                            "content": full_response
                        }
                    )
                    rag_chat_text = ""

                    for msg in st.session_state.rag_messages:
                        rag_chat_text += (
                            f"{msg['role']} : "
                            f"{msg['content']}\n\n"
                        )
                    messages_json = json.dumps(
                        st.session_state.rag_messages,
                        ensure_ascii=False
                    )
                    if (
                            st.session_state.get(
                                "current_rag_chat_id"
                            ) is None
                    ):

                        rag_chat_id = save_chat(
                            f"RAG - {rag_prompt[:30]}",
                            rag_chat_text,
                            messages_json,
                            st.session_state.userid
                        )

                        st.session_state.current_rag_chat_id = (
                            rag_chat_id
                        )

                    else:

                        update_chat(
                            st.session_state.current_rag_chat_id,
                            f"RAG - {rag_prompt[:30]}",
                            rag_chat_text,
                            messages_json,
                            st.session_state.userid
                        )



                    if "selected_document" not in st.session_state:
                        st.session_state.selected_document = None
                    results = search_chunks(
                        rag_prompt,
                        selected_document=
                        st.session_state.selected_document,
                        top_k=3,
                        user_id=st.session_state.userid
                    )
                    with st.expander(
                            "🔍 Similarity Results"
                    ):

                        for score, chunk in results:
                            st.write(
                                f"""
                    Document:
                    {chunk[1]}
                    
                    Chunk:
                    {chunk[0]}

                    Similarity:
                    {score:.4f}
                    """
                            )

                            st.code(
                                chunk[2][:500]
                            )
                    if results:
                        best_score, best_chunk = results[0]

                        st.info(
                            f"""
                        📄 Source Document:
                        {best_chunk[1]}
                        
                        CHunk Id:
                        {best_chunk[0]}


                        ⭐ Similarity:
                        {best_score:.4f}
                        """
                        )

                        st.rerun()




            except Exception as e:

                st.error(
                    f"Ollama Error: {e}"
                )